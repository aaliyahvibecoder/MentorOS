import streamlit as st
import utils
import datetime
import pandas as pd
import json
from pypdf import PdfReader
import docx

st.set_page_config(page_title="New Client | WSO OS", layout="wide")
utils.load_css()
utils.init_db()

st.title("CLIENT INTAKE PROTOCOL")

# --- TABBED INTERFACE ---
tab1, tab2 = st.tabs(["👤 INDIVIDUAL INTAKE", "📦 BULK CSV IMPORT"])

# =========================================================
# TAB 1: INDIVIDUAL INTAKE (Standard Wizard)
# =========================================================
with tab1:
    st.info("ENTER DETAILS TO GENERATE NEW DOSSIER AND SYNC RESUME DATA.")

    # --- SECTION 1: LOGISTICS (UPDATED WITH DATE & REMINDERS) ---
    with st.container(border=True):
        st.markdown("### 1. IDENTITY & SCHEDULING")
        c1, c2, c3 = st.columns(3)
        
        # 1. NAME (Column 1)
        with c1:
            client_name = st.text_input("STUDENT NAME (Required)*")
        
        # 2. SESSION TYPE (Column 3) 
        # Executed BEFORE Column 2 so we can check the type for Async logic
        with c3:
            session_type = st.selectbox("INITIAL SESSION TYPE", [
                "Mock Interview", 
                "Career Roadmap", 
                "7 Stories Review", 
                "Resume Review (Full)", 
                "LinkedIn Audit", 
                "Networking Strategy"
            ])
            # Reminder Freq
            reminder = st.selectbox("REMINDER FREQUENCY", ["None", "24h Before", "1h Before"])

        # 3. DATE & TIME (Column 2)
        with c2:
            session_date = st.date_input("NEXT SESSION DATE", datetime.date.today())
            
            # LOGIC: If Resume Review, hide time input (Async)
            if session_type == "Resume Review (Full)":
                st.info("🕒 Time: **Async** (No fixed slot)")
                session_time = None
            else:
                session_time = st.time_input("SESSION TIME", datetime.time(9, 00))

    # --- SECTION 2: RESUME DATA SYNC ---
    with st.container(border=True):
        st.markdown("### 2. RESUME CONFIGURATION")
        st.caption("Upload here to sync the CV directly to the database for future Resume Review sessions.")
        
        rc1, rc2 = st.columns([1, 2])
        with rc1:
            is_exp = st.checkbox("EXPERIENCED HIRE? (Deal Exp > Education)", help="Check if candidate has significant deal experience. Defaults to Undergrad format if unchecked.")
        with rc2:
            cv_file = st.file_uploader("UPLOAD CV (PDF/DOCX) - AUTO-PARSE", type=["pdf", "docx"])
            
        # --- ROBUST & DEDUPLICATED TEXT EXTRACTOR ---
        extracted_text = ""
        if cv_file:
            try:
                # 1. Handle PDF
                if cv_file.name.lower().endswith('.pdf'):
                    reader = PdfReader(cv_file)
                    for page in reader.pages:
                        extracted_text += page.extract_text() or ""
                
                # 2. Handle DOCX (with Merged Cell Deduplication)
                elif cv_file.name.lower().endswith('.docx'):
                    doc = docx.Document(cv_file)
                    full_text = []
                    
                    # A. Grab Standard Paragraphs
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    
                    # B. Grab Text from Tables (Handling Merged Cells)
                    for table in doc.tables:
                        for row in table.rows:
                            # Unique set to track cells we've already read in this row
                            seen_cells = set()
                            for cell in row.cells:
                                if cell not in seen_cells:
                                    seen_cells.add(cell)
                                    for para in cell.paragraphs:
                                        full_text.append(para.text)
                                    
                    extracted_text = "\n".join(full_text)
                
                if len(extracted_text) > 50:
                    st.success(f"✅ CV Parsed Successfully ({len(extracted_text)} chars). Ready to sync.")
                else:
                    st.warning("⚠️ File uploaded, but very little text found. Is it an image scan?")
                    
            except Exception as e:
                st.error(f"Error parsing file: {e}")

    # --- SECTION 3: ASSESSMENT ---
    with st.container(border=True):
        st.markdown("### 3. PRE-SESSION ASSESSMENT")
        
        objective = st.text_input("PRIMARY OBJECTIVE", placeholder="e.g., Secure 1st round at Goldman Sachs...")
        
        c_str, c_wk = st.columns(2)
        with c_str:
            strengths = st.text_area("PERCEIVED STRENGTHS", height=100, placeholder="e.g., High GPA, unique military background, clean resume.")
        with c_wk:
            weaknesses = st.text_area("PERCEIVED WEAKNESSES", height=100, placeholder="e.g., Tends to ramble, weak on DCF mechanics, no target schools.")

    st.markdown("---")

    # --- EXECUTION ENGINE ---
    if st.button("INITIALIZE CLIENT DOSSIER", use_container_width=True):
        if client_name.strip():
            upper_name = client_name.strip().upper()
            
            # Format Date/Time
            if session_time:
                time_str = session_time.strftime("%H:%M")
            else:
                time_str = "Async"
            
            date_str = session_date.strftime("%Y-%m-%d")
                
            str_val = strengths if strengths else "Pending Assessment"
            foc_val = weaknesses if weaknesses else "Pending Assessment"
            hist_val = "New Client Onboarding Completed."
            exp_int = 1 if is_exp else 0 
            
            conn = utils.get_db_connection()
            cursor = conn.cursor()
            
            # UPDATED INSERT: Added session_date and reminder_freq
            cursor.execute('''
                INSERT INTO clients (student, session_date, time, reminder_freq, type, strengths, focus, history, resume_text, is_experienced)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (upper_name, date_str, time_str, reminder, session_type, str_val, foc_val, hist_val, extracted_text, exp_int))
            conn.commit()

            # Refresh Session State from DB to ensure new columns are picked up
            cursor.execute("SELECT * FROM clients")
            st.session_state['client_db'] = [dict(row) for row in cursor.fetchall()]
            
            st.success(f"DOSSIER GENERATED: {upper_name} scheduled for {date_str} @ {time_str}")
            st.toast(f"Intake Complete: {upper_name}")
        else:
            st.error("ERROR: STUDENT NAME IS REQUIRED TO INITIALIZE A DOSSIER.")

# =========================================================
# TAB 2: BULK IMPORT (CSV Logic)
# =========================================================
with tab2:
    st.info("PROTOCOL: UPLOAD A CSV WITH THE HEADERS BELOW TO IMPORT MULTIPLE CLIENTS AT ONCE.")
    st.markdown("#### 1. REQUIRED CSV FORMAT")
    st.caption("Headers: Student Name, Session Type, Strengths, Weaknesses, History, Mock Score (Tech), Mock Score (Beh)")
    
    bulk_file = st.file_uploader("Upload Client List (CSV)", type=["csv"])
    
    if bulk_file:
        try:
            df = pd.read_csv(bulk_file)
            st.write("Preview of Upload:")
            st.dataframe(df.head(), use_container_width=True)
            
            if st.button("🚀 EXECUTE BULK IMPORT", type="primary", use_container_width=True):
                success_count = 0
                conn = utils.get_db_connection()
                cursor = conn.cursor()
                progress_bar = st.progress(0)
                total_rows = len(df)
                
                # Bulk import generally doesn't set specific future dates, defaulting to today or None if needed
                today_str = datetime.date.today().strftime("%Y-%m-%d")

                for index, row in df.iterrows():
                    name = str(row.get("Student Name", "Unknown")).strip().upper()
                    sType = str(row.get("Session Type", "General"))
                    strs = str(row.get("Strengths", "Imported"))
                    wks = str(row.get("Weaknesses", "Imported"))
                    hist = str(row.get("History", "Imported via CSV."))
                    default_time = "Async" if "Resume" in sType else "09:00"

                    mock_json = "[]"
                    try:
                        t_score = float(row.get("Mock Score (Tech)", 0))
                        b_score = float(row.get("Mock Score (Beh)", 0))
                        if t_score > 0 or b_score > 0:
                            mock_entry = [{
                                "date": today_str,
                                "tech": t_score,
                                "beh": b_score,
                                "notes": "Imported Historical Data"
                            }]
                            mock_json = json.dumps(mock_entry)
                    except:
                        mock_json = "[]"

                    try:
                        cursor.execute('''
                            INSERT INTO clients (student, session_date, time, type, strengths, focus, history, mock_data, resume_text, is_experienced)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (name, today_str, default_time, sType, strs, wks, hist, mock_json, "", 0))
                        success_count += 1
                    except Exception as e:
                        st.warning(f"Skipped row {index}: {e}")
                    
                    progress_bar.progress((index + 1) / total_rows)

                conn.commit()
                cursor.execute("SELECT * FROM clients")
                rows = cursor.fetchall()
                st.session_state['client_db'] = [dict(row) for row in rows]
                
                st.success(f"✅ SUCCESS: {success_count} CLIENTS IMPORTED.")
                st.balloons()
                
        except Exception as e:
            st.error(f"Error reading CSV: {e}")
