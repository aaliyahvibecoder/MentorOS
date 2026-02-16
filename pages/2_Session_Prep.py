import streamlit as st
import utils
import datetime
import json
import io
import csv
from pypdf import PdfReader
import docx

st.set_page_config(page_title="Session Prep | WSO OS", layout="wide")
utils.load_css()
utils.init_db()

st.title("SESSION EXECUTION")

# =========================================================
# LAYER 0: GLOBAL MASTER BANK (SIDEBAR)
# =========================================================
with st.sidebar:
    st.markdown("### 🏛️ MASTER KNOWLEDGE VAULT")
    st.caption("Documents uploaded here are globally accessible for all sessions.")
    
    master_upload = st.file_uploader("Add to Master Bank", type=["pdf", "docx", "txt", "csv"], key="master_up")
    if master_upload:
        m_content = ""
        try:
            if master_upload.name.lower().endswith('.pdf'):
                reader = PdfReader(master_upload)
                m_content = "\n".join([p.extract_text() for p in reader.pages])
            elif master_upload.name.lower().endswith('.docx'):
                doc = docx.Document(master_upload)
                full_text = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        seen = set()
                        for cell in row.cells:
                            if cell not in seen:
                                seen.add(cell)
                                full_text.extend([p.text for p in cell.paragraphs])
                m_content = "\n".join(full_text)
            elif master_upload.name.lower().endswith('.csv'):
                m_content = master_upload.getvalue().decode("utf-8")
            elif master_upload.name.lower().endswith('.txt'):
                m_content = str(master_upload.read(), "utf-8")

            if st.button("💾 SAVE TO MASTER VAULT"):
                conn = utils.get_db_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO global_kb (filename, content) VALUES (?, ?)", (master_upload.name, m_content))
                conn.commit()
                st.success(f"Archived: {master_upload.name}")
                st.rerun()
        except Exception as e:
            st.error(f"Upload Error: {e}")

    if st.session_state.get('global_kb'):
        st.markdown("---")
        st.markdown("**CURRENT VAULT INVENTORY:**")
        for doc in st.session_state['global_kb']:
            st.caption(f"• {doc['filename']}")
        
        if st.button("🗑️ WIPE MASTER VAULT", type="secondary"):
            conn = utils.get_db_connection()
            conn.cursor().execute("DELETE FROM global_kb")
            conn.commit()
            st.rerun()

# =========================================================
# MAIN INTERFACE: SELECTION & CONTEXT
# =========================================================
client_names = [c['student'] for c in st.session_state['client_db']]
client_list = ["Guest / Walk-in"] + client_names

col_sel1, col_sel2 = st.columns([1, 2])
with col_sel1:
    selected_client = st.selectbox("SELECT CANDIDATE", client_list)

client_kb_text = ""
if selected_client != "Guest / Walk-in":
    client_data = next((c for c in st.session_state['client_db'] if c['student'] == selected_client), None)
    if client_data:
        client_kb_text = client_data.get('session_kb_text') or ""

with col_sel2:
    session_type = st.selectbox("SELECT SESSION TYPE", [
        "Mock Interview", "Career Roadmap", "LinkedIn Audit", "7 Stories Review", "Networking Strategy"
    ])

st.markdown("---")

# =========================================================
# MOCK INTERVIEW MODULE (FULL PERSISTENCE + REPORT CARD)
# =========================================================
if session_type == "Mock Interview":
    st.header(f"MOCK INTERVIEW: {selected_client}")
    
    with st.expander("📂 CLIENT-SPECIFIC OVERLAY (JOB DESCS / NOTES)", expanded=False):
        st.markdown("Upload documents unique to **this specific student**.")
        client_files = st.file_uploader("Upload to Client Dossier", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        
        if client_files:
            new_client_text = ""
            for cf in client_files:
                if cf.name.lower().endswith('.pdf'):
                    reader = PdfReader(cf)
                    new_client_text += "\n".join([p.extract_text() for p in reader.pages])
                elif cf.name.lower().endswith('.docx'):
                    doc = docx.Document(cf)
                    new_client_text += "\n".join([p.text for p in doc.paragraphs])
            
            if st.button("💾 SAVE TO CLIENT DOSSIER"):
                updated_kb = (client_kb_text + "\n" + new_client_text).strip()
                conn = utils.get_db_connection()
                conn.cursor().execute("UPDATE clients SET session_kb_text = ? WHERE student = ?", (updated_kb, selected_client))
                conn.commit()
                st.success("Dossier Updated.")
                st.rerun()

    # --- THE UNIFIED INTERVIEW ENGINE ---
    with st.container(border=True):
        st.subheader("🤖 INTERVIEW ENGINE (TECHNICAL & BEHAVIORAL)")
        
        c1, c2 = st.columns(2)
        with c1:
            tech_topics = st.multiselect("TECHNICAL FOCUS", ["Accounting", "Valuation", "DCF", "LBO", "M&A", "Markets", "RX"])
        with c2:
            beh_topics = st.multiselect("BEHAVIORAL FOCUS", ["TMAY", "Why IB/PE?", "Conflict/Teamwork", "Strengths/Weaknesses", "Market View"])
        
        use_manual = st.toggle("⚡ ENABLE AI PROMPT OVERRIDE")
        if use_manual:
            custom_prompt = st.text_area("ENTER CUSTOM AI INSTRUCTIONS", placeholder="e.g., Focus on the deal experience mentioned in their CV...")
        
        if st.button("🚀 GENERATE INTEGRATED SESSION PLAN", use_container_width=True):
            with st.spinner("Analyzing Master Vault and Dossier..."):
                try:
                    engine = utils.AIEngine()
                    master_context = "\n\n".join([d['content'] for d in st.session_state.get('global_kb', [])])
                    
                    system_instruction = f"""
                    You are a strict Wall Street Managing Director. 
                    MASTER KNOWLEDGE BASE: {master_context[:20000]}
                    CLIENT-SPECIFIC DOSSIER: {client_kb_text[:5000]}
                    """
                    
                    if use_manual:
                        final_prompt = f"{system_instruction}\n\nUSER OVERRIDE INSTRUCTION: {custom_prompt}"
                    else:
                        final_prompt = f"""
                        {system_instruction}
                        TASK:
                        1. Create a structured 45-minute agenda.
                        2. Generate Technical Questions for: {tech_topics}.
                        3. Generate Behavioral Questions for: {beh_topics}.
                        4. For EACH question, provide the Question, 1-2 Follow-ups, and the Target WSO Answer.
                        Format: Markdown with Bold headers.
                        """
                    
                    response = engine.generate_content(final_prompt)
                    # STORE IN SESSION STATE IMMEDIATELY
                    st.session_state[f"last_agenda_{selected_client}"] = response.text
                except Exception as e:
                    st.error(f"AI Error: {e}")

        # DISPLAY PERSISTENTLY
        if f"last_agenda_{selected_client}" in st.session_state:
             st.markdown("### 📋 INTEGRATED SESSION PLAN")
             st.info(st.session_state[f"last_agenda_{selected_client}"])

    # --- SCORING & LOGGING ---
    st.markdown("---")
    sc1, sc2 = st.columns(2)
    tech_score = sc1.slider("TECHNICAL SCORE", 1, 10, 5)
    beh_score = sc2.slider("BEHAVIORAL SCORE", 1, 10, 5)
    live_notes = st.text_area("LIVE SESSION NOTES & FEEDBACK")

    c_save, c_pdf = st.columns(2)
    
    with c_save:
        if st.button("💾 SAVE SESSION & LOG QUESTIONS", use_container_width=True):
            if selected_client == "Guest / Walk-in":
                st.error("Cannot save guest.")
            else:
                # FIX: Explicitly grab full text from session state to avoid truncation
                agenda = st.session_state.get(f"last_agenda_{selected_client}", "No AI Agenda generated.")
                timestamp = datetime.date.today().strftime("%m/%d")
                
                # Formatting the log entry to include everything
                full_log = f"""
                SESSION: Mock Interview ({tech_score} T, {beh_score} B)
                FEEDBACK: {live_notes}
                AI AGENDA / QUESTIONS ASKED:
                {agenda}
                """
                
                # Use a separator for the history column
                history_update = f" | {timestamp}: {full_log}"
                
                new_mock = {"date": timestamp, "tech": tech_score, "beh": beh_score, "notes": live_notes}

                conn = utils.get_db_connection()
                cursor = conn.cursor()
                cursor.execute("SELECT history, mock_data FROM clients WHERE student = ?", (selected_client,))
                row = cursor.fetchone()
                
                if row:
                    m_data = json.loads(row['mock_data']) if row['mock_data'] else []
                    m_data.append(new_mock)
                    cursor.execute("UPDATE clients SET history = ?, mock_data = ? WHERE student = ?", 
                                   ((row['history'] or "") + history_update, json.dumps(m_data), selected_client))
                    conn.commit()
                    # Sync State
                    for c in st.session_state['client_db']:
                        if c['student'] == selected_client:
                            c['history'] = (row['history'] or "") + history_update
                            c['mock_data'] = m_data
                            break
                    st.success("✅ FULL SESSION SAVED. Context stored for future reference.")
    
    with c_pdf:
        # --- NEW: PDF REPORT CARD GENERATOR ---
        if st.button("📄 GENERATE REPORT CARD (PDF)", use_container_width=True):
            if selected_client == "Guest / Walk-in":
                st.error("Please select a valid client from DB.")
            else:
                try:
                    pdf_buffer = utils.create_pdf_report(
                        student=selected_client,
                        tech_score=tech_score,
                        beh_score=beh_score,
                        feedback=live_notes,
                        agenda=st.session_state.get(f"last_agenda_{selected_client}", "")
                    )
                    
                    st.download_button(
                        label="⬇️ DOWNLOAD PDF SCORECARD",
                        data=pdf_buffer,
                        file_name=f"WSO_Report_{selected_client}.pdf",
                        mime="application/pdf",
                        type="primary"
                    )
                except Exception as e:
                    st.error(f"Error generating PDF: {e}")

# =========================================================
# OTHER SESSION TYPES (Standard Logic)
# =========================================================
elif session_type == "7 Stories Review":
    st.header(f"7 STORIES REVIEW: {selected_client}")
    st.info("PROTOCOL: 60 MIN SESSION. ENSURE STAR FRAMEWORK ADHERENCE.")
    
    story_prompts = ["1. Strengths", "2. Weaknesses", "3. Hard Worker", "4. Team Player", "5. Conflict (Peer)", "6. Conflict (Superior)", "7. Ethical"]
    current_stories = {}
    progress_bar = st.progress(0)
    
    for i, story in enumerate(story_prompts):
        clean_key = f"{selected_client}_{story.replace(' ', '_')}"
        with st.expander(story.upper(), expanded=False):
            c1, c2 = st.columns([3, 1])
            note = c1.text_area("STORY NOTES (S.T.A.R.)", height=150, key=f"note_{clean_key}", placeholder="Situation... Task... Action... Result...")
            c2.markdown("**VERIFICATION**")
            star_checked = c2.checkbox("✅ STAR VERIFIED", key=f"star_{clean_key}")
            
            if c2.button("🤖 AI STAR AUDIT", key=f"ai_btn_{clean_key}", use_container_width=True):
                if not note or len(note) < 10:
                    c2.error("Need more notes.")
                else:
                    with st.spinner("Scanning..."):
                        try:
                            model = utils.AIEngine()
                            system_prompt = """
                            Analyze these notes against the STAR framework.
                            Return valid JSON with keys: "Situation", "Task", "Action", "Result" (values: "[Found/Missing] - Brief reason"), and "Verdict" ("Pass"/"Needs Polish").
                            """
                            response = model.generate_content(f"{system_prompt}\n\nNOTES:\n{note}", config={"response_mime_type": "application/json"})
                            audit_data = json.loads(response.text)
                            v_color = "green" if audit_data.get("Verdict") == "Pass" else "red"
                            st.markdown(f"**VERDICT:** :{v_color}[{audit_data.get('Verdict')}]")
                            st.json(audit_data)
                        except Exception as e: st.error(f"Audit Failed: {e}")
            
            current_stories[story] = {"notes": note, "star_verified": star_checked}

    verified_count = sum(1 for s in current_stories.values() if s['star_verified'])
    progress_bar.progress(verified_count / 7.0)
    st.caption(f"**{verified_count} / 7 Stories Verified**")
    st.markdown("---")
    
    if st.button("SAVE STORY PROTOCOL"):
        if selected_client == "Guest / Walk-in": st.warning("Cannot save guest.")
        else:
            timestamp = datetime.date.today().strftime("%m/%d")
            history_entry = f" | {timestamp}: 7 Stories Review ({verified_count}/7 STAR verified)"
            stories_json = json.dumps(current_stories)
            conn = utils.get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT history FROM clients WHERE student = ?", (selected_client,))
            row = cursor.fetchone()
            if row:
                new_history = row['history'] or ""
                if history_entry not in new_history: new_history += history_entry
                cursor.execute("UPDATE clients SET history = ?, stories_log = ? WHERE student = ?", (new_history, stories_json, selected_client))
                conn.commit()
                for client in st.session_state['client_db']:
                    if client['student'] == selected_client:
                        client['stories_log'] = current_stories
                        client['history'] = new_history
                        break
                st.success(f"SAVED TO DATABASE: {selected_client}")
            else: st.error("Client not found.")

elif session_type == "LinkedIn Audit":
    import datetime
    st.header(f"LINKEDIN AUDIT: {selected_client}")
    c_head, c_about = st.columns(2)
    with c_head:
        headline = st.text_input("PASTE HEADLINE:")
        hard_skills = ["modeling", "valuation", "excel", "finance", "accounting", "m&a", "analyst"]
        has_hard_skills = any(skill in headline.lower() for skill in hard_skills) if headline else False
    with c_about:
        about_section = st.text_area("PASTE ABOUT SECTION:")
        first_person_triggers = [" i ", " i'm ", " my ", " me ", "i've"]
        padded_about = f" {about_section.lower()} "
        is_first_person = any(fp in padded_about for fp in first_person_triggers) if about_section else False

    st.markdown("### 📊 WSO RULE CHECKLIST")
    c1, c2 = st.columns(2)
    with c1:
        photo_check = st.checkbox("1. PHOTO: BUSINESS FORMAL", help="Visual check required.")
        headline_check = st.checkbox("2. HEADLINE: HARD SKILLS", value=has_hard_skills)
    with c2:
        about_check = st.checkbox("3. ABOUT: 1ST PERSON ONLY", value=is_first_person)
        url_check = st.checkbox("4. CUSTOM URL SET")
        
    if st.button("SAVE LINKEDIN PROTOCOL"):
        if selected_client == "Guest / Walk-in": st.error("Cannot save guest.")
        else:
            timestamp = datetime.date.today().strftime("%m/%d")
            score = sum([photo_check, headline_check, about_check, url_check])
            history_entry = f" | {timestamp}: LinkedIn Audit ({score}/4 Rules Passed)"
            conn = utils.get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT history FROM clients WHERE student = ?", (selected_client,))
            row = cursor.fetchone()
            if row:
                new_history = (row['history'] or "") + history_entry
                cursor.execute("UPDATE clients SET history = ? WHERE student = ?", (new_history, selected_client))
                conn.commit()
                for client in st.session_state['client_db']:
                    if client['student'] == selected_client:
                        client['history'] = new_history
                        break
                st.success(f"SAVED: {selected_client} scored {score}/4")

elif session_type == "Career Roadmap":
    import datetime
    st.header(f"CAREER ROADMAP: {selected_client}")
    
    st.markdown("### 1. CANDIDATE CONTEXT")
    c1, c2 = st.columns(2)
    with c1:
        road_cv = st.file_uploader("UPLOAD CANDIDATE CV (PDF)", type=["pdf"], key="road_cv")
        cv_text = ""
        if road_cv:
            try:
                reader = PdfReader(road_cv)
                for page in reader.pages: cv_text += page.extract_text() or ""
                st.success("✅ CV Parsed.")
            except: st.error("Error reading CV.")
    with c2:
        aspirations = st.text_area("EXPERIENCES & ASPIRATIONS", height=150, placeholder="e.g., Target: Mega-fund PE...")

    st.markdown("### 2. FIRM TARGETING")
    t1, t2, t3 = st.columns(3)
    reach = t1.text_area("REACH FIRMS")
    core = t2.text_area("CORE FIRMS")
    safety = t3.text_area("SAFETY FIRMS")

    st.markdown("### 3. CRITICAL HURDLES")
    chk1, chk2, chk3, chk4 = st.columns(4)
    chk_visa = chk1.checkbox("VISA REQUIRED?")
    chk_timeline = chk2.checkbox("OFF-CYCLE?")
    chk_gpa = chk3.checkbox("GPA < 3.5")
    chk_prestige = chk4.checkbox("NON-TARGET")

    if st.button("🤖 GENERATE STRATEGY SUMMARY", use_container_width=True):
        if not aspirations and not reach: st.warning("Enter notes/targets.")
        else:
            with st.spinner("Synthesizing..."):
                try:
                    model = utils.AIEngine()
                    hurdles = []
                    if chk_visa: hurdles.append("Visa Sponsorship")
                    if chk_timeline: hurdles.append("Off-cycle")
                    if chk_gpa: hurdles.append("Low GPA")
                    if chk_prestige: hurdles.append("Non-target")
                    
                    system_prompt = f"""
                    Synthesize this student data into a 1-paragraph Career Strategy Summary.
                    Raw CV: {cv_text[:1500]}...
                    Aspirations: {aspirations}
                    Targets: Reach: {reach}, Core: {core}, Safety: {safety}
                    Hurdles: {', '.join(hurdles)}
                    """
                    response = model.generate_content(system_prompt)
                    st.markdown("#### 🎯 STRATEGY SUMMARY")
                    st.info(response.text)
                    st.session_state[f'roadmap_summary_{selected_client}'] = response.text
                except Exception as e: st.error(f"AI Error: {e}")

    if st.button("SAVE ROADMAP TO DOSSIER"):
        if selected_client == "Guest / Walk-in": st.error("Cannot save guest.")
        else:
            timestamp = datetime.date.today().strftime("%m/%d")
            history_entry = f" | {timestamp}: Career Roadmap Completed."
            conn = utils.get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT history FROM clients WHERE student = ?", (selected_client,))
            row = cursor.fetchone()
            if row:
                new_history = (row['history'] or "") + history_entry
                summary_key = f'roadmap_summary_{selected_client}'
                if summary_key in st.session_state: new_history += f" (Strategy: {st.session_state[summary_key][:50]}...)"
                cursor.execute("UPDATE clients SET history = ? WHERE student = ?", (new_history, selected_client))
                conn.commit()
                for client in st.session_state['client_db']:
                    if client['student'] == selected_client:
                        client['history'] = new_history
                        break
                st.success(f"SAVED: {selected_client}")

elif session_type == "Networking Strategy":
    st.header(f"NETWORKING STRATEGY: {selected_client}")
    email_draft = st.text_area("PASTE STUDENT EMAIL DRAFT HERE:", height=250)
    
    word_count = len(email_draft.split()) if email_draft else 0
    is_concise = 0 < word_count <= 100
    has_time = any(kw in email_draft.lower() for kw in ["monday", "tuesday", "available", "time", "chat"]) if email_draft else False
    has_cg = any(kw in email_draft.lower() for kw in ["alumni", "university", "background", "both"]) if email_draft else False

    c1, c2, c3 = st.columns(3)
    c1.metric("WORD COUNT (<100)", word_count)
    c1.checkbox("1. COMMON GROUND?", value=has_cg)
    c2.metric("CONCISE?", "PASS" if is_concise else "FAIL")
    c2.checkbox("2. SPECIFIC TIMES?", value=has_time)
    c3.metric("STATUS", "ACTIVE" if email_draft else "IDLE")
    c3.checkbox("3. CONCISE?", value=is_concise)

    mentor_notes = st.text_area("MENTOR NOTES", height=100)

    if st.button("SAVE NETWORKING PROTOCOL"):
        if selected_client == "Guest / Walk-in": st.error("Cannot save guest.")
        else:
            timestamp = datetime.date.today().strftime("%m/%d")
            score = sum([is_concise, has_time, has_cg]) # Simplified scoring
            history_entry = f" | {timestamp}: Networking Strategy ({score}/3 Rules Passed)"
            conn = utils.get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT history FROM clients WHERE student = ?", (selected_client,))
            row = cursor.fetchone()
            if row:
                new_history = (row['history'] or "") + history_entry
                cursor.execute("UPDATE clients SET history = ? WHERE student = ?", (new_history, selected_client))
                conn.commit()
                for client in st.session_state['client_db']:
                    if client['student'] == selected_client:
                        client['history'] = new_history
                        break
                st.success(f"SAVED: {selected_client}")
